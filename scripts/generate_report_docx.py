import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: Markdown file not found at {md_path}")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Simple Markdown Parser
    in_math_block = False
    math_buffer = []

    for line in lines:
        line = line.strip()
        
        # Skip empty lines (except inside math block?)
        if not line and not in_math_block:
            continue
            
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
            
        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            # Ordered list
            content = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(content, style='List Number')
            
        # Horizontal Rule
        elif line.startswith('---'):
            doc.add_page_break()
            
        # LaTeX Math Block handling
        elif line.startswith('$$'):
            if not in_math_block:
                in_math_block = True
                math_buffer = []
                # Check if it's a single line block like $$ x=1 $$
                if len(line) > 2 and line.endswith('$$'):
                     in_math_block = False
                     math_text = line.strip('$')
                     p = doc.add_paragraph()
                     run = p.add_run(math_text)
                     run.italic = True
                     run.font.color.rgb = RGBColor(0, 0, 128)
                     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                in_math_block = False
                math_text = '\n'.join(math_buffer)
                p = doc.add_paragraph()
                run = p.add_run(math_text)
                run.italic = True
                run.font.color.rgb = RGBColor(0, 0, 128)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif in_math_block:
            math_buffer.append(line)
            
        # Normal text with bold/italic parsing
        else:
            p = doc.add_paragraph()
            # Handle inline math $...$
            # Simple regex for bold **text** and math $...$
            # This is a naive parser, splitting by tokens
            parts = re.split(r'(\*\*.*?\*\*|\$.*?\$)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('$') and part.endswith('$'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                    run.font.color.rgb = RGBColor(0, 0, 128)
                else:
                    p.add_run(part)

    try:
        doc.save(docx_path)
        print(f"Successfully generated Word document at: {docx_path}")
    except Exception as e:
        print(f"Failed to save document: {e}")

if __name__ == "__main__":
    md_file = os.path.join("reports", "Project_Progress_Report.md")
    docx_file = os.path.join("reports", "Project_Progress_Report.docx")
    markdown_to_docx(md_file, docx_file)
